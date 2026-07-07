import docx
import glob
import os

def extract_docx_text(file_path):
    doc = docx.Document(file_path)
    fullText = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            fullText.append(para.text.strip())
            
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not row_text or row_text[-1] != cell_text:
                    row_text.append(cell_text)
            if row_text:
                fullText.append(" | ".join(row_text))
                
    return "\n".join(fullText)

if __name__ == '__main__':
    paths = [
        "/Users/jax/Documents/Obowiązakowe informacje CLP na etykiecie Płyn do.docx",
        "/Users/jax/Documents/Obowiązakowe informacje CLP na etykiecie.docx"
    ]
    
    for path in paths:
        if os.path.exists(path):
            name = os.path.basename(path).replace(" ", "_")
            try:
                txt = extract_docx_text(path)
                out_path = f"/Users/jax/Desktop/plyndo.pl-main/extracted_{name}.txt"
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(txt)
                print(f"Success: {path} -> {out_path}")
            except Exception as e:
                print(f"Error {path}: {e}")
        else:
            print(f"File not found: {path}")
